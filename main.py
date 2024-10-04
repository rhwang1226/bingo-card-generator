from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from random import sample
from docx.shared import Inches

def generate_bingo_card():
    bingo_card = []
    for i in range(5):
        if i == 0:
            col_numbers = sample(range(1, 11), 5)
        elif i == 1:
            col_numbers = sample(range(11, 21), 5)
        elif i == 2:
            col_numbers = sample(range(21, 31), 5)
        elif i == 3:
            col_numbers = sample(range(31, 41), 5)
        elif i == 4:
            col_numbers = sample(range(41, 51), 5)
        col_numbers.sort()
        bingo_card.append(col_numbers)
    
    print(bingo_card)
    return list(map(list, zip(*bingo_card)))
    

def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

doc = Document()

section = doc.sections[-1]
section.page_width = Inches(11)
section.page_height = Inches(8.5)

section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

for page_num in range(200):
    title = doc.add_paragraph("SASE Trivia Bingo")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)

    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in table.rows:
        row.height = Inches(1.3)

    bingo_card = generate_bingo_card()

    for i in range(5):
        for j in range(5):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(str(bingo_card[i][j]))
            run.font.size = Pt(11)
            set_borders(cell)

    if page_num < 199:
        doc.add_page_break()

doc.save("SASE-Bingo-Cards.docx")

print("Document created successfully!")